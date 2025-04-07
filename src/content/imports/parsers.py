from django.utils.html import format_html
from django.utils.safestring import SafeText, mark_safe
from docx.drawing import Drawing
from docx.image.image import Image
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.rel import _Relationship
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from wagtail_content_import.parsers.microsoft import DocxParser
from wagtail_content_import.parsers.tables import Cell, Table

HEADINGS_MAPPING: dict[str, str] = {
    "Heading 1": "1",
    "Heading 2": "2",
    "Heading 3": "3",
    "Heading 4": "4",
    "Heading 5": "5",
    "Heading 1_DBT": "1",
    "Heading 2_DBT": "2",
    "Heading 3_DBT": "3",
    "Heading 4_DBT": "4",
    "Heading 5_DBT": "5",
}
LIST_MAPPING: dict[str, str] = {
    "DBT num list": "ol",
    "Bullet List 1": "ul",
    "List Paragraph": "ul",
}
PARAGRAPH_NAMES: list[str] = [
    "Normal",
]

KNOWN_NAMES: list[str] = (
    list(HEADINGS_MAPPING.keys()) + PARAGRAPH_NAMES + list(LIST_MAPPING.keys())
)


class DocxParser(DocxParser):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        self.images = [
            rel for _, rel in self.document.part.rels.items() if rel.reltype == RT.IMAGE
        ]
        self.tables = [
            Table(
                rows=[
                    [Cell(cell.text) for cell in row.cells]
                    for row in document_table.rows
                ]
            )
            for document_table in self.document.tables
        ]

    def generate_a_tag(self, content, href):
        return (
            format_html('<a href="{href}">{content}</a>', href=href, content=content)
            if content
            else ""
        )

    def paragraph_to_heading(
        self, paragraph: Paragraph, heading_level
    ) -> None | dict[str, str]:
        if not paragraph.text:
            return None

        block_val = paragraph.text.strip()
        if not block_val:
            return None

        return {
            "type": f"heading{heading_level}",
            "value": block_val,
        }

    def paragraph_to_html(self, paragraph: Paragraph, block_type, block_data):
        """
        Compile a paragraph into a HTML string, optionally with semantic markup for styles.
        Returns a dictionary of the form:
        {
            'type': 'html',
            'value': html
        }
        OR
        {
            'type': 'html-list',
            'list_type': 'ol' | 'ul',
            'value': html
        }
        """
        outer_tag = "p"

        text_list = []
        for content in paragraph.iter_inner_content():
            text = content.text
            match content:
                case Run():
                    if text:
                        if content.bold:
                            text = self.generate_simple_tag(text, "b")
                        if content.italic:
                            text = self.generate_simple_tag(text, "em")
                        text_list.append(text)
                        continue

                    for item in content.iter_inner_content():
                        if isinstance(item, Drawing):
                            image_rel: _Relationship = self.images.pop(0)
                            if not image_rel.is_external:
                                image = Image.from_blob(image_rel.target_part.blob)
                                self.blocks.append({"type": "image", "value": image})
                                continue
                            else:
                                raise Exception("External images not yet supported")

                case Hyperlink():
                    text_list.append(self.generate_a_tag(text, content.address))
                case _:
                    raise Exception(f"Unknown content type: {type(content)}")

        if not text_list:
            return None

        content: SafeText = mark_safe("".join(text_list))  # noqa: S308

        block = {
            "type": "html",
            "value": "",
        }

        if block_type == "list":
            outer_tag = "li"
            block = {
                "type": "html-list",
                "list_type": block_data["list_type"],
                "indent_level": paragraph.style.paragraph_format.left_indent,
            }

        block["value"] = self.generate_simple_tag(content, outer_tag)
        return block

    def get_block_type_for_paragraph(self, paragraph: Paragraph) -> tuple[str, dict]:
        if heading_level := HEADINGS_MAPPING.get(paragraph.style.name, None):
            return "heading", {
                "heading_level": heading_level,
            }

        if list_type := LIST_MAPPING.get(paragraph.style.name, None):
            return "list", {
                "list_type": list_type,
                "list_indent": 0,
            }

        return "html", {}

    def parse(self):
        """
        Parse the document and return a set of intermediate {'type': type, 'value': value} blocks that represent it.
        """

        title = self.document.core_properties.title

        self.blocks = []

        paragraphs: list[Paragraph] = self.document.paragraphs
        for paragraph in paragraphs:
            block_type, block_data = self.get_block_type_for_paragraph(paragraph)
            if block_type == "heading":
                heading_level = block_data["heading_level"]

                if heading_level == "1" and not title:
                    title = paragraph.text.strip()

                if heading_block := self.paragraph_to_heading(paragraph, heading_level):
                    self.blocks.append(heading_block)
                continue

            converted_block = self.paragraph_to_html(
                paragraph,
                block_type,
                block_data,
            )
            if converted_block:
                self.blocks.append(converted_block)

        # SECOND PASS

        final_blocks = []
        html_content = ""

        for i, block in enumerate(self.blocks):
            prev_block = self.blocks[i - 1] if i > 0 else None
            next_block = self.blocks[i + 1] if i < len(self.blocks) - 1 else None

            if block["type"] in ["html", "html-list"]:
                # Open the list
                if block["type"] == "html-list":
                    if (
                        not prev_block
                        or prev_block["type"] != "html-list"
                        or prev_block["list_type"] != block["list_type"]
                    ):
                        # Open the list
                        html_content += f"<{block['list_type']}>"

                # Add the html content
                html_content += block["value"]

                # Close the list
                if block["type"] == "html-list":
                    if (
                        not next_block
                        or next_block["type"] != "html-list"
                        or next_block["list_type"] != block["list_type"]
                    ):
                        # Close the list
                        html_content += f"</{block['list_type']}>"

                if not next_block or next_block["type"] not in ["html", "html-list"]:
                    final_blocks.append(
                        {
                            "type": "html",
                            "value": html_content,
                        }
                    )
                    html_content = ""
            else:
                final_blocks.append(block)

        # Add tables to final blocks
        for table in self.tables:
            final_blocks.append(
                {
                    "type": "table",
                    "value": table,
                }
            )

        # # Add images to final blocks
        # for document_image in self.document.inline_shapes:
        #     print(f"need to generate an image for : {type(document_image)}")

        return {"title": title, "elements": final_blocks}
